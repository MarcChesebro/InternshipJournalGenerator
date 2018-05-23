import datetime
from dateutil import parser
import re
from docx import Document
from docx.shared import Pt

TAGS = ['bold', 'size', '/bold', '/size']


def load_template():
    with open('_template.txt', 'r') as content_file:
        content = content_file.read()
    return content


def write_file(name, content):
    with open(name, 'w+') as file:
        file.write(content)


def add_week_days(week_date, date_list):
    for i in range(0, 5):
        date_list.append(week_date + datetime.timedelta(days=i))


def create_date_list(date):
    dates = [date]
    # journal date

    week1_date = date - datetime.timedelta(days=14)
    week2_date = date - datetime.timedelta(days=7)
    end_date = date - datetime.timedelta(days=3)

    dates.append(week1_date)
    dates.append(end_date)

    dates.append(week1_date)
    add_week_days(week1_date, dates)

    dates.append(week2_date)
    add_week_days(week2_date, dates)

    return dates


def format_journal(dates, template):
    datestrings = [date.strftime("%m/%d/%Y") for date in dates]
    return template.format(*datestrings)


def parse_tags(template):
    lines = [re.split(r'[<|>]', line) for line in template.splitlines()]
    for line in lines:
        line.append('\n')
    return lines


def docx_conversion(journal_lines, input_document):
    # not too sure how this should look
    paragraph = input_document.add_paragraph()
    _is_bold_ = False
    _is_size_ = False
    for line in journal_lines:
        for e in line:
            if e in TAGS:
                _is_bold_, _is_size_ = process_tag(e, _is_bold_, _is_size_)
            else:
                run = paragraph.add_run(e)
                run.bold = _is_bold_
                if _is_size_:
                    font = run.font
                    font.size = Pt(14)


def process_tag(input_tag, bold, size):
    if input_tag == TAGS[0]:
        bold = True
    elif input_tag == TAGS[1]:
        size = True
    elif input_tag == TAGS[2]:
        bold = False
    elif input_tag == TAGS[3]:
        size = False
    return bold, size


def generate_journal(date):
    template = load_template()
    dates = create_date_list(date)
    file_name = '{}-InternshipJournal-MarcChesebro.docx'.format(date.strftime("%m-%d-%Y"))
    journal_string = format_journal(dates, template)
    journal_lines = parse_tags(journal_string)
    document = Document()
    docx_conversion(journal_lines, document)
    document.save(file_name)


if __name__ == '__main__':
    datestring = input('Enter the due date: ')
    date = parser.parse(datestring)

    generate_journal(date)
    print("done")
